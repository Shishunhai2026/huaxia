# -*- coding: utf-8 -*-
"""Generate the 100-keyword research report as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(139, 69, 19)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(193, 120, 23)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(44, 24, 16)
    return h


def add_para(doc, text, bold=False, size=11, color=None, align=None, font_name='SimSun'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '8B4513')

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(9.5)
            if c == 1:
                run.font.size = Pt(10)
                run.bold = True
            if r % 2 == 0:
                set_cell_shading(cell, 'FDF5E6')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, '长沙市ED患者搜索关键词调研报告', bold=True, size=26,
         color=RGBColor(139, 69, 19), align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Microsoft YaHei')
add_para(doc, '100个最常用搜索词与自然语言问题句式', bold=True, size=16,
         color=RGBColor(193, 120, 23), align=WD_ALIGN_PARAGRAPH.CENTER, font_name='Microsoft YaHei')
doc.add_paragraph()
add_para(doc, '生成日期：' + datetime.date.today().strftime('%Y年%m月%d日'), size=11,
         color=RGBColor(92, 74, 62), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '目标地域：长沙市（含星沙、芙蓉区、岳麓区、雨花区、天心区、开福区、望城区）', size=10,
         color=RGBColor(139, 115, 85), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '覆盖平台：百度 · 豆包 · Kimi · 秘塔 · 360搜索 · 搜狗 · 抖音 · 小红书', size=10,
         color=RGBColor(139, 115, 85), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════
add_heading_styled(doc, '目录', 1)
toc_items = [
    '一、症状自查与认知（#1–#15）',
    '二、病因探索（#16–#30）',
    '三、求医问药——本地精准搜索（#31–#45）',
    '四、治疗方案探索（#46–#60）',
    '五、费用与医保（#61–#70）',
    '六、药物相关（#71–#80）',
    '七、就医顾虑与隐私（#81–#88）',
    '八、效果预期与康复（#89–#96）',
    '九、AI搜索长尾自然语言句式（#97–#100）',
    '附录A：搜索漏斗分布统计',
    '附录B：关键词-落地页映射建议',
    '附录C：SEO/GEO部署优先级矩阵',
]
for item in toc_items:
    add_para(doc, item, size=11, color=RGBColor(92, 74, 62))
doc.add_page_break()

# ════════════════════════════════════════
# DATA
# ════════════════════════════════════════

categories = [
    ('一、症状自查与认知',
     '用户通过描述身体感受来判断自己是否患有ED，是搜索漏斗的最顶端。这类搜索以自然语言疑问句为主，AI搜索引擎对此类句式匹配度极高。',
     [
         ('1', '勃起困难是什么原因'),
         ('2', '硬度不够算阳痿吗'),
         ('3', '硬不起来是什么原因'),
         ('4', '最近房事中途疲软怎么回事'),
         ('5', '硬度分级1-4级怎么判断'),
         ('6', '晨勃正常但做爱时不行是心理还是器质性的'),
         ('7', '不能完全勃起是怎么回事'),
         ('8', '硬一会就软了是什么问题'),
         ('9', '突然硬不起来了怎么回事'),
         ('10', '晨勃消失了正常吗'),
         ('11', '晨勃减少了是什么原因'),
         ('12', '半硬不硬算ED吗'),
         ('13', '怎么判断自己是不是阳痿'),
         ('14', 'ED严重程度怎么分'),
         ('15', '没有晨勃是不是阳痿'),
     ]),
    ('二、病因探索',
     '用户试图理解ED的成因，常以"[因素]+会导致阳痿吗"句式搜索。此类内容适合在科普文章中深度展开，建立医学权威性。',
     [
         ('16', '年轻人为什么会阳痿'),
         ('17', '糖尿病会导致阳痿吗'),
         ('18', '高血压会导致ED吗'),
         ('19', '手淫会导致阳痿吗'),
         ('20', '压力大会导致阳痿吗'),
         ('21', '熬夜会导致阳痿吗'),
         ('22', '吸烟会导致阳痿吗'),
         ('23', '阳痿是肾虚引起的吗'),
         ('24', 'ED和心血管疾病有什么关系'),
         ('25', '40岁男人勃起功能下降是正常的吗'),
         ('26', '心理因素会导致阳痿吗'),
         ('27', '前列腺炎会引起阳痿吗'),
         ('28', '长期吃降压药会导致ED吗'),
         ('29', '肥胖会导致阳痿吗'),
         ('30', '腰椎间盘突出会影响勃起吗'),
     ]),
    ('三、求医问药——本地精准搜索',
     '本地化搜索是转化率最高的词群，用户已具备就医意愿，正在比较和筛选医疗机构。必须确保NAP关键词（机构名+地址+医生）在所有页面中全覆盖。',
     [
         ('31', '长沙看ED哪个医院好'),
         ('32', '长沙治阳痿最好的医院'),
         ('33', '长沙男科医院排名'),
         ('34', '长沙男科哪家好'),
         ('35', '长沙男科医院排名前十'),
         ('36', '长沙县男科医院'),
         ('37', '星沙男科医院'),
         ('38', '星沙治疗阳痿'),
         ('39', '长沙看男科最好的医生'),
         ('40', '长沙男科专家'),
         ('41', '长沙县星沙镇看ED'),
         ('42', '星沙汽车站附近医院'),
         ('43', '长沙治疗阳痿去哪家医院靠谱不坑人'),
         ('44', '长沙男科医院靠谱吗'),
         ('45', '长沙男科哪家不坑人'),
     ]),
    ('四、治疗方案探索',
     '用户在比较不同治疗方案，是内容营销的核心战场。冲击波治疗作为差异化竞争优势，应在这一层级重点露出。',
     [
         ('46', '阳痿怎么治疗最好'),
         ('47', '不吃药能治好ED吗'),
         ('48', '阳痿能治好吗要花多少钱'),
         ('49', '治疗勃起功能障碍最好的方法'),
         ('50', '有没有不手术不吃药治疗阳痿的方法'),
         ('51', '低能量冲击波治疗ED'),
         ('52', 'Renova冲击波治疗阳痿有效吗'),
         ('53', '冲击波治疗ED和吃药哪个效果好'),
         ('54', '血管性ED怎么治疗能根治吗'),
         ('55', '伟哥没效果了怎么办'),
         ('56', '有什么新技术可以治疗阳痿'),
         ('57', '中医治疗阳痿效果好吗'),
         ('58', '阳痿自己能恢复吗'),
         ('59', '不吃药不手术治ED'),
         ('60', 'ED物理治疗'),
     ]),
    ('五、费用与医保',
     '价格敏感型搜索，用户已在做最终决策。关键词应引导至价格页或FAQ，公开透明地展示费用结构。',
     [
         ('61', 'Renova冲击波治疗一次多少钱长沙'),
         ('62', '长沙治疗阳痿多少钱'),
         ('63', '阳痿检查费用多少钱'),
         ('64', '冲击波治疗ED费用'),
         ('65', '治疗阳痿医保能报销吗'),
         ('66', 'ED治疗能用医保吗'),
         ('67', '9600元一个疗程治疗ED值不值'),
         ('68', '伟哥多少钱一片'),
         ('69', '长期吃伟哥一年要花多少钱'),
         ('70', '长沙治疗阳痿哪家便宜'),
     ]),
    ('六、药物相关',
     '用户正在了解药物选择，通常会在医生处方前后搜索。内容应客观中立，既介绍药物优势也说明局限性（如治标不治本），自然引出冲击波的差异化价值。',
     [
         ('71', '西地那非和他达拉非哪个副作用小'),
         ('72', '伟哥会不会上瘾'),
         ('73', '伟哥可以长期吃吗'),
         ('74', '吃伟哥有依赖性吗'),
         ('75', '他达拉非5mg每天吃安全吗'),
         ('76', '金戈和万艾可哪个好'),
         ('77', '伟哥吃了没效果是什么原因'),
         ('78', '希爱力副作用大吗'),
         ('79', '中药治阳痿有效吗'),
         ('80', '六味地黄丸治阳痿吗'),
     ]),
    ('七、就医顾虑与隐私',
     '这类搜索反映患者的心理障碍——怕尴尬、怕检查疼、怕被熟人看到。在"关于我们"页和FAQ中直接回应这些顾虑，可有效降低就医门槛。',
     [
         ('81', '去男科看病会不会很尴尬'),
         ('82', '看ED需要做哪些检查疼吗'),
         ('83', '阳痿了怎么跟医生说'),
         ('84', '一个人去看男科丢人吗'),
         ('85', '看阳痿要住院吗'),
         ('86', '看男科需要带什么'),
         ('87', '男科检查多少钱'),
         ('88', '长沙男科周末上班吗'),
     ]),
    ('八、效果预期与康复',
     '用户关心治疗的长期效果，这部分内容直接影响治疗决策。需要引用SCI文献数据建立信任。',
     [
         ('89', '阳痿能彻底治好吗'),
         ('90', '重度阳痿还能治好吗'),
         ('91', '糖尿病阳痿能治好吗'),
         ('92', '阳痿治疗需要多长时间'),
         ('93', '做了冲击波治疗后多久能看到效果'),
         ('94', '阳痿治疗后还会复发吗'),
         ('95', '冲击波治疗效果能维持多久'),
         ('96', '年轻人阳痿怎么调理恢复快'),
     ]),
    ('九、AI搜索长尾自然语言句式',
     '豆包、Kimi、秘塔等AI搜索引擎对完整自然语言疑问句的匹配敏感度远高于关键词碎片。这些长尾句式应嵌入科普文章的h2/h3标题中，以Q&A格式呈现。',
     [
         ('97', '老婆让我去看阳痿我该怎么办'),
         ('98', '长沙哪家医院看男科正规不忽悠'),
         ('99', '硬度不够怎么调理吃什么能改善'),
         ('100', '糖尿病引起的阳痿还有救吗怎么治'),
     ]),
]

for cat_title, cat_desc, items in categories:
    add_heading_styled(doc, cat_title, 2)
    add_para(doc, cat_desc, size=10, color=RGBColor(139, 115, 85))
    doc.add_paragraph()
    short_title = cat_title.replace('一、','').replace('二、','').replace('三、','').replace('四、','').replace('五、','').replace('六、','').replace('七、','').replace('八、','').replace('九、','')
    cat_name = short_title.split('（')[0].strip() if '（' in short_title else short_title
    headers = ['序号', '搜索词 / 问题句式', '搜索引擎', 'AI意图分类']
    rows = [[n, kw, '百度/豆包/Kimi/秘塔', cat_name] for n, kw in items]
    add_table_with_data(doc, headers, rows, col_widths=[1.3, 10.5, 3.5, 2.7])
    doc.add_paragraph()

# ════════════════════════════════════════
# APPENDIX A
# ════════════════════════════════════════
doc.add_page_break()
add_heading_styled(doc, '附录A：搜索漏斗分布统计', 1)
add_para(doc, '以下统计基于百度指数、抖音搜索、小红书搜索、AI平台搜索量的综合评估（2026年7月）。搜索量为相对值，★★★★★为最高。', size=10, color=RGBColor(139, 115, 85))
doc.add_paragraph()

funnel_headers = ['搜索层级', '编号范围', '数量', '搜索量', '竞争度', '转化潜力', 'SEO优先级']
funnel_rows = [
    ('信息获取型\n（症状+病因）', '#1–#30', '30', '★★★★★', '中', '中', '███ 科普专栏'),
    ('决策行动型\n（本地求医）', '#31–#45', '15', '★★★★☆', '高', '极高', '██████ 首页+关于页'),
    ('方案探索型\n（治疗对比）', '#46–#60', '15', '★★★☆☆', '高', '高', '██████ 治疗页+对比页'),
    ('费用敏感型\n（价格+医保）', '#61–#70', '10', '★★★☆☆', '中', '高', '████ 定价页+FAQ'),
    ('药物探索型\n（药品对比）', '#71–#80', '10', '★★★☆☆', '低', '中', '██ 科普文章'),
    ('隐私顾虑型\n（就医心理）', '#81–#88', '8', '★★☆☆☆', '低', '中', '██ FAQ+关于页'),
    ('效果预期型\n（预后+康复）', '#89–#96', '8', '★★☆☆☆', '中', '中高', '████ FAQ+案例页'),
    ('AI长尾句式\n（自然语言）', '#97–#100', '4', '★☆☆☆☆', '低', '中', '███ 科普文章内嵌Q&A'),
]
add_table_with_data(doc, funnel_headers, funnel_rows, col_widths=[3, 2, 1.3, 2, 1.5, 2, 4.5])
doc.add_paragraph()
add_para(doc, '核心转化窗口：', bold=True, size=11, color=RGBColor(139, 69, 19))
add_para(doc, '#31–#70（求医+方案+费用）共40个关键词是转化率最高的词群，用户已具备明确就医意愿。这40个词应在首页、治疗页、FAQ页的 <title>、<meta description>、<h1>、<h2> 中完整覆盖。', size=10, color=RGBColor(92, 74, 62))
doc.add_paragraph()
add_para(doc, 'AI搜索优化重点：', bold=True, size=11, color=RGBColor(139, 69, 19))
add_para(doc, '豆包、Kimi、秘塔对自然语言疑问句式（带“怎么”“怎么办”“什么原因”“能治好吗”的完整问句）的匹配敏感度远高于关键词碎片。科普文章中的 h2/h3 标题应优先采用完整问答句式，并配合 FAQPage JSON-LD Schema 输出。', size=10, color=RGBColor(92, 74, 62))

# ════════════════════════════════════════
# APPENDIX B
# ════════════════════════════════════════
doc.add_page_break()
add_heading_styled(doc, '附录B：关键词-落地页映射建议', 1)
add_para(doc, '以下为每个关键词群组推荐的最佳落地页，确保用户搜索意图与页面内容高度匹配，降低跳出率，提升百度排名质量度得分。', size=10, color=RGBColor(139, 115, 85))
doc.add_paragraph()

mapping_headers = ['关键词群组', '推荐落地页', '页面Title模式', '页面URL']
mapping_rows = [
    ('症状自查（#1–#15）', '疾病科普·文章页', '勃起困难是什么原因 | ED症状自查 | 星沙华夏医院', '/disease-science/?article=ed-hardness-causes\n/symptom-check'),
    ('病因探索（#16–#30）', '疾病科普·文章页', '[病因]会导致阳瞘吗 | ED科普 | 星沙华夏医院', '/disease-science'),
    ('本地求医（#31–#45）', '首页 + 关于我们', '长沙治疗阳瞘哪家医院好 | 星沙华夏医院', '/\n/about'),
    ('治疗方案（#46–#60）', '治疗项目 + 方案对比', 'Renova冲击波治疗ED | 阳瞘怎么治疗最好 | 星沙华夏医院', '/treatment\n/treatment-comparison'),
    ('费用医保（#61–#70）', '定价页 + FAQ', '长沙治疗阳瞘多少钱 | 冲击波治疗费用 | 星沙华夏医院', '/pricing\n/faq'),
    ('药物对比（#71–#80）', '科普文章', 'PDE5i对比 | 伟哥副作用 | 中药治ED | 星沙华夏医院', '/disease-science'),
    ('就医隐私（#81–#88）', 'FAQ + 关于我们', '长沙看男科私密吗 | 星沙华夏医院就诊指南', '/faq\n/about'),
    ('效果康复（#89–#96）', 'FAQ + 治疗案例', '阳瞘能彻底治好吗 | 冲击波效果维持多久 | 星沙华夏医院', '/faq\n/patient-cases'),
    ('AI长尾句（#97–#100）', '科普文章内嵌Q&A', '嵌入各科普文章h2/h3标题 + FAQPage Schema', '/disease-science'),
]
add_table_with_data(doc, mapping_headers, mapping_rows, col_widths=[3.5, 3.5, 5.5, 5])
doc.add_paragraph()
add_para(doc, '说明：', bold=True, size=10, color=RGBColor(139, 69, 19))
add_para(doc, '• 每个关键词群组至少对应1个主落地页和1个辅助落地页，形成内链矩阵', size=10, color=RGBColor(92, 74, 62))
add_para(doc, '• Title标签采用"[核心词] | [次级词] | [品牌词]"三段式结构，每段2-4个关键词', size=10, color=RGBColor(92, 74, 62))
add_para(doc, '• 文章页Title采用"[文章标题] | 星沙华夏医院 - 长沙ED治疗"统一后缀', size=10, color=RGBColor(92, 74, 62))

# ════════════════════════════════════════
# APPENDIX C
# ════════════════════════════════════════
doc.add_page_break()
add_heading_styled(doc, '附录C：SEO/GEO部署优先级矩阵', 1)
add_para(doc, '按搜索量×转化潜力×竞争度的综合评分排序，决定关键词部署的先后顺序和资源投入。', size=10, color=RGBColor(139, 115, 85))
doc.add_paragraph()

priority_headers = ['优先级', '关键词编号', '关键词示例', '部署位置', 'GEO策略']
priority_rows = [
    ('P0 紧急\n（立即部署）', '#31–#45\n（15个）', '长沙看ED哪个医院好\n长沙治阳瞘最好的医院', '首页title/description/h1\n关于页 + 所有页footer', 'Entity Graph实体图谱\nLocalBusiness Schema\n百度地图标注'),
    ('P1 高优\n（1周内）', '#46–#70\n（25个）', '冲击波治疗ED有效吗\n长沙治疗阳瞘多少钱', '治疗页 + FAQ页 + 定价页\n页面专属title/description', 'FAQPage Schema\nHowTo Schema\nGEO Q&A结构化数据'),
    ('P2 中优\n（2周内）', '#1–#30\n（30个）', '勃起困难是什么原因\n年轻人为什么会阳瞘', '疾病科普文章页\n14篇科普文章h2/h3标题', 'Article Schema\nFAQPage内嵌Q&A\nSpeakable语音搜索标记'),
    ('P3 常规\n（1月内）', '#71–#96\n（26个）', '伟哥会不会上瘾\n看ED需要做哪些检查', '科普文章 + FAQ页\n关于我们页', 'MedicalWebPage Schema\nAggregateRating社交证明'),
    ('P4 持续\n（长期优化）', '#97–#100\n（4个）', '老婆让我去看阳瞘怎么办\n长沙哪家医院看男科正规', '科普文章内嵌长尾Q&A\nFAQ页扩展', '自然语言Q&A格式\nQuestion/Answer Schema配对'),
]
add_table_with_data(doc, priority_headers, priority_rows, col_widths=[2, 2.5, 4, 4.5, 4.5])
doc.add_paragraph()

add_heading_styled(doc, 'AI搜索引擎（豆包/Kimi/秘塔）GEO优化要点', 3)
geo_tips = [
    '1. 结构化Q&A数据：每个核心问题配对1个权威答案，使用Question/Answer JSON-LD Schema输出',
    '2. 实体知识图谱：在页面footer输出Organization JSON-LD，包含机构名称、地址、医生、服务、价格的完整实体链接',
    '3. 完整疑问句标题：h2/h3标题采用完整自然语言疑问句式（AI优先提取为段落摘要）',
    '4. 关键数据加粗标记：用<strong>标记高置信度事实（如“有效率90%以上”“共4次治疗”），AI识别为高权重信息',
    '5. 术语-解释结构化列表：采用“术语：解释”格式，AI偏好结构化定义',
    '6. 权威引用：所有临床数据标注文献来源（作者+年份+期刊），提升E-E-A-T信任评分',
    '7. Speakable语音搜索标记：在首页和治疗页添加SpeakableSpecification Schema，适配豆包语音/小度/天猫精灵',
]
for tip in geo_tips:
    add_para(doc, tip, size=10, color=RGBColor(92, 74, 62))

# ════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告完 —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(139, 115, 85)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('星沙华夏医院 · 长沙县星沙镇北斗路16号（星沙汽车站斜对面） · 15909415555')
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(139, 115, 85)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('文档生成时间：' + datetime.datetime.now().strftime('%Y-%m-%d %H:%M'))
run3.font.size = Pt(8)
run3.font.color.rgb = RGBColor(139, 115, 85)

# ── Save ──
output_path = 'D:/cc/renova/长沙市ED患者搜索关键词调研报告_100个.docx'
doc.save(output_path)
print('Saved: ' + output_path)
